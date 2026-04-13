from pathlib import Path

import pytest
from docx import Document

from lexai.pipeline.arguments import IRACBlock, IssueArguments
from lexai.pipeline.issues import Issue
from lexai.pipeline.memorial import MemorialInput, build_memorial
from lexai.pipeline.research import RankedCase


def _sample_input() -> MemorialInput:
    issues = [
        Issue(
            issue_title="Whether CCL has jurisdiction over patented products",
            area_of_law="competition",
            relevant_statutes=["Competition Act 2002", "Patents Act 1970"],
            relevant_articles=[],
            description="Whether the exercise of patent rights ousts CCI jurisdiction.",
        ),
    ]
    cases = [
        RankedCase(
            case_slug="2020-scc-online-del-598",
            case_title="Monsanto Holdings v CCI",
            citation="2020 SCC OnLine Del 598",
            court="Delhi High Court",
            year=2020,
            area_of_law="competition",
            best_chunk_text="holding patent rights and competition jurisdiction are not mutually exclusive",
            reasoning="Monsanto directly addresses patent-CCI overlap",
        ),
    ]
    args = [
        IssueArguments(
            petitioner_arguments=[
                IRACBlock(
                    issue="Jurisdiction over patented products",
                    rule="Section 3/4 + Section 60 Competition Act",
                    application="Per 2020 SCC OnLine Del 598, patent rights do not oust Section 3/4.",
                    conclusion="CCL has jurisdiction.",
                )
            ],
            respondent_arguments=[
                IRACBlock(
                    issue="Patent-linked conduct excluded",
                    rule="Patents Act primacy",
                    application="The impugned conduct is patent-rights exercise.",
                    conclusion="CCL lacks jurisdiction.",
                )
            ],
        )
    ]
    return MemorialInput(
        moot_title="NALSAR-CCI Anti-Trust Moot Court 2026 — Sample",
        tribunal="Competition Appellate Tribunal of Lilliput",
        case_number="Appeal No. __ of 2026",
        issues=issues,
        arguments_per_issue=args,
        cases_per_issue=[cases],
    )


def test_build_memorial_writes_valid_docx(tmp_path: Path):
    out = tmp_path / "memorial.docx"
    path = build_memorial(_sample_input(), out)
    assert path == out
    assert out.exists()
    assert out.stat().st_size > 1000

    # open it — this proves the file is a readable docx, not garbage
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "NALSAR" in text
    assert "Arguments Advanced" in text
    assert "Jurisdiction over patented products" in text or "patented products" in text
    assert "Prayer" in text


def test_build_memorial_table_of_authorities_lists_cited_cases(tmp_path: Path):
    out = tmp_path / "memorial.docx"
    build_memorial(_sample_input(), out)
    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    assert "Monsanto Holdings v CCI" in full_text
    assert "2020 SCC OnLine Del 598" in full_text


def test_build_memorial_raises_on_mismatched_input_lengths(tmp_path: Path):
    inp = _sample_input()
    inp.arguments_per_issue.append(inp.arguments_per_issue[0])  # one more args than issues
    with pytest.raises(ValueError):
        build_memorial(inp, tmp_path / "m.docx")
