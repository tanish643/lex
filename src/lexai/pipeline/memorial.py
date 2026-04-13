"""DOCX memorial assembly.

Stitches issues + arguments + cited cases into a structured moot-court
memorial. Sections match Indian moot conventions:

  1. Cover Page
  2. Table of Contents (manual stub — Word auto-fills on open)
  3. Index of Authorities (all cited cases, deduplicated)
  4. Statement of Jurisdiction
  5. Statement of Facts (placeholder — human fills in)
  6. Issues Raised
  7. Summary of Arguments
  8. Arguments Advanced (IRAC per issue, both sides)
  9. Prayer

Formatting is intentionally minimal. A moot team will re-style in Word;
the product promises content quality, not typography.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pydantic import BaseModel

from lexai.pipeline.arguments import IssueArguments
from lexai.pipeline.issues import Issue
from lexai.pipeline.research import RankedCase


class MemorialInput(BaseModel):
    moot_title: str
    tribunal: str
    case_number: str
    issues: list[Issue]
    arguments_per_issue: list[IssueArguments]
    cases_per_issue: list[list[RankedCase]]


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def _add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _cover_page(doc: Document, inp: MemorialInput) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(inp.moot_title.upper())
    run.bold = True
    run.font.size = Pt(18)

    for _ in range(2):
        doc.add_paragraph()
    tb = doc.add_paragraph()
    tb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tb.add_run(f"BEFORE THE\n{inp.tribunal.upper()}")
    run.bold = True
    run.font.size = Pt(14)

    for _ in range(2):
        doc.add_paragraph()
    case = doc.add_paragraph()
    case.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = case.add_run(inp.case_number)
    run.font.size = Pt(12)

    doc.add_page_break()


def _table_of_authorities(doc: Document, inp: MemorialInput) -> None:
    _add_heading(doc, "Index of Authorities", level=1)

    # dedupe cases across issues by case_slug
    unique: dict[str, RankedCase] = {}
    for issue_cases in inp.cases_per_issue:
        for c in issue_cases:
            unique.setdefault(c.case_slug, c)

    if not unique:
        _add_para(doc, "(no authorities cited)")
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Case"
    hdr[1].text = "Citation"
    hdr[2].text = "Court"
    for c in sorted(unique.values(), key=lambda c: c.year):
        row = table.add_row().cells
        row[0].text = c.case_title
        row[1].text = c.citation
        row[2].text = f"{c.court} ({c.year})"
    doc.add_page_break()


def _statement_of_jurisdiction(doc: Document, inp: MemorialInput) -> None:
    _add_heading(doc, "Statement of Jurisdiction", level=1)
    _add_para(
        doc,
        f"The Appellants have approached the {inp.tribunal} under the "
        "relevant appellate provisions of the Competition Act, 2002. "
        "The Respondents respectfully submit to the jurisdiction of this "
        "Tribunal for the present proceedings.",
    )
    doc.add_page_break()


def _statement_of_facts(doc: Document) -> None:
    _add_heading(doc, "Statement of Facts", level=1)
    _add_para(
        doc,
        "[Placeholder — the moot team should insert a concise narrative of "
        "facts here based on the problem. AI cannot generate a trustworthy "
        "summary of facts because factual misstatements are fatal in court.]",
    )
    doc.add_page_break()


def _issues_raised(doc: Document, inp: MemorialInput) -> None:
    _add_heading(doc, "Issues Raised", level=1)
    for i, issue in enumerate(inp.issues, start=1):
        _add_para(doc, f"{i}. {issue.issue_title}", bold=True)
    doc.add_page_break()


def _summary_of_arguments(doc: Document, inp: MemorialInput) -> None:
    _add_heading(doc, "Summary of Arguments", level=1)
    for i, (issue, args) in enumerate(
        zip(inp.issues, inp.arguments_per_issue), start=1
    ):
        _add_heading(doc, f"Issue {i}: {issue.issue_title}", level=2)
        if args.petitioner_arguments:
            _add_para(doc, "Appellant contends:", bold=True)
            _add_para(doc, args.petitioner_arguments[0].conclusion)
        if args.respondent_arguments:
            _add_para(doc, "Respondent contends:", bold=True)
            _add_para(doc, args.respondent_arguments[0].conclusion)
    doc.add_page_break()


def _arguments_advanced(doc: Document, inp: MemorialInput) -> None:
    _add_heading(doc, "Arguments Advanced", level=1)
    for i, (issue, args, cases) in enumerate(
        zip(inp.issues, inp.arguments_per_issue, inp.cases_per_issue), start=1
    ):
        _add_heading(doc, f"Issue {i}: {issue.issue_title}", level=2)

        _add_heading(doc, "Arguments on behalf of the Appellant", level=3)
        for block in args.petitioner_arguments:
            _add_para(doc, f"Issue: {block.issue}")
            _add_para(doc, f"Rule: {block.rule}")
            _add_para(doc, f"Application: {block.application}")
            _add_para(doc, f"Conclusion: {block.conclusion}")
            doc.add_paragraph()

        _add_heading(doc, "Arguments on behalf of the Respondent", level=3)
        for block in args.respondent_arguments:
            _add_para(doc, f"Issue: {block.issue}")
            _add_para(doc, f"Rule: {block.rule}")
            _add_para(doc, f"Application: {block.application}")
            _add_para(doc, f"Conclusion: {block.conclusion}")
            doc.add_paragraph()

        if cases:
            _add_para(doc, "Authorities relied upon:", bold=True)
            for c in cases:
                _add_para(doc, f"  • {c.case_title}, {c.citation}")
        doc.add_page_break()


def _prayer(doc: Document) -> None:
    _add_heading(doc, "Prayer", level=1)
    _add_para(
        doc,
        "In light of the facts stated, issues raised, arguments advanced, "
        "and authorities cited, the Honourable Tribunal may graciously be "
        "pleased to adjudge and declare in favour of the submitting party; "
        "and pass such further orders as this Tribunal may deem fit in the "
        "interests of justice, equity, and good conscience.",
    )
    _add_para(doc, "")
    _add_para(doc, "All of which is respectfully submitted.", bold=True)


def build_memorial(inp: MemorialInput, out_path: Path) -> Path:
    if len(inp.issues) != len(inp.arguments_per_issue):
        raise ValueError(
            f"issues ({len(inp.issues)}) and arguments_per_issue "
            f"({len(inp.arguments_per_issue)}) length mismatch"
        )
    if len(inp.issues) != len(inp.cases_per_issue):
        raise ValueError(
            f"issues ({len(inp.issues)}) and cases_per_issue "
            f"({len(inp.cases_per_issue)}) length mismatch"
        )

    doc = Document()

    _cover_page(doc, inp)
    _table_of_authorities(doc, inp)
    _statement_of_jurisdiction(doc, inp)
    _statement_of_facts(doc)
    _issues_raised(doc, inp)
    _summary_of_arguments(doc, inp)
    _arguments_advanced(doc, inp)
    _prayer(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
